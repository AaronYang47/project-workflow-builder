from __future__ import annotations

from collections import Counter
import hashlib
from pathlib import Path
import re

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "artifacts" / "PROFAB_Project_Lifecycle_JF_Traceability_Verification_2026-08-30.docx"
FORM_SOURCE = ROOT / "src" / "lib" / "profab-forms.ts"
SCREENSHOTS = ROOT / "artifacts" / "screenshots"
QUOTE_ASSETS = ROOT / "tmp" / "report-quote-images"
CJK_FONT_PATH = Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf")

NAVY = "0B1F33"
BLUE = "1F5A94"
TEAL = "159A85"
GREEN = "0F7B5A"
AMBER = "B85D05"
RED = "B42318"
VIOLET = "6653A6"
INK = "172033"
MUTED = "5D6A7E"
LINE = "D7E0EA"
PALE_BLUE = "EEF5FB"
PALE_TEAL = "EAF7F3"
PALE_AMBER = "FFF6E8"
PALE_RED = "FDEEEE"
PALE_GREY = "F5F7FA"
WHITE = "FFFFFF"

FONT_LATIN = "Aptos"
# Use a CJK family shipped with macOS so the mixed-language source quotations
# remain visible in both Word and the LibreOffice render used for final QA.
FONT_CJK = "Arial Unicode MS"


NODE_LABELS = {
    "project-start": "Project Start",
    "gate-g1-qualified": "G1 — Qualified & Commercially Engaged",
    "pre-construction": "Pre-Construction",
    "gate-g2-technical-commitment": "G2 — Project / Technical Commitment",
    "production-readiness": "Production Readiness",
    "gate-g3-production-authorization": "G3 — Production Authorization",
    "factory-production": "Factory Production",
    "gate-g4-factory-release": "G4 — Factory Completion / Release",
    "delivery-project-completion": "Delivery / Project Completion",
    "gate-g5-warranty-start": "G5 — Project Completion / Warranty Start",
    "commissioning-warranty": "Commissioning & Warranty",
    "close-out": "Final Close",
}

SHORT_NODE_LABELS = {
    "project-start": "Start",
    "gate-g1-qualified": "G1",
    "pre-construction": "Pre-Con",
    "gate-g2-technical-commitment": "G2",
    "production-readiness": "Prod. Ready",
    "gate-g3-production-authorization": "G3",
    "factory-production": "Factory",
    "gate-g4-factory-release": "G4",
    "delivery-project-completion": "Delivery",
    "gate-g5-warranty-start": "G5",
    "commissioning-warranty": "Warranty",
    "close-out": "Final Close",
}


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, size: float | None = None, bold: bool | None = None,
                 color: str | None = None, italic: bool | None = None,
                 font: str = FONT_LATIN):
    run.font.name = font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:eastAsia"), FONT_CJK)
    if font == FONT_CJK:
        r_fonts.set(qn("w:hint"), "eastAsia")
        lang = r_pr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            r_pr.append(lang)
        lang.set(qn("w:eastAsia"), "zh-CN")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = rgb(color)
    if italic is not None:
        run.italic = italic
    return run


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=95, bottom=80, end=95):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_fixed(table):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    table.autofit = False


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def cell_paragraph(cell, text: str, size=8.0, color=INK, bold=False,
                   italic=False, space_after=0, alignment=None):
    paragraph = cell.paragraphs[0] if not cell.text and len(cell.paragraphs) == 1 else cell.add_paragraph()
    if paragraph.text:
        paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.02
    if alignment is not None:
        paragraph.alignment = alignment
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return paragraph


def set_cell_lines(cell, lines, size=8.0):
    cell.text = ""
    for index, line in enumerate(lines):
        is_image = isinstance(line, dict) and bool(line.get("image_path"))
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1.5 if index < len(lines) - 1 else 0)
        paragraph.paragraph_format.line_spacing = 1.02
        if is_image:
            run = paragraph.add_run()
            picture = run.add_picture(
                str(line["image_path"]),
                width=Inches(line.get("width_inches", 4.62)),
            )
            alt_text = line.get("alt_text", "")
            if alt_text:
                picture._inline.docPr.set("descr", alt_text)
                picture._inline.docPr.set("title", "JF source wording")
            continue
        if isinstance(line, str):
            text, bold, color, italic = line, False, INK, False
        else:
            text = line.get("text", "")
            bold = line.get("bold", False)
            color = line.get("color", INK)
            italic = line.get("italic", False)
            font = line.get("font", FONT_LATIN)
        if isinstance(line, str):
            font = FONT_LATIN
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color, italic=italic, font=font)


def _text_width(draw, text, font):
    if not text:
        return 0
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0]


def _wrap_mixed_language_text(draw, text, font, max_width):
    lines = []
    for source_line in text.splitlines() or [""]:
        tokens = re.findall(r"[\u3400-\u9fff]|[^\u3400-\u9fff\s]+|\s+", source_line)
        current = ""
        for token in tokens:
            token = " " if token.isspace() else token
            candidate = current + token
            if not current or _text_width(draw, candidate, font) <= max_width:
                current = candidate
                continue
            lines.append(current.rstrip())
            current = token.lstrip()
        lines.append(current.rstrip())
    return lines or [""]


def build_source_quote_image(text, key, width_px=1380, font_size_px=29):
    """Rasterize mixed Chinese/English source wording for reliable Word/PDF display.

    The exact wording is also stored as picture alt text in the DOCX. This avoids
    a known LibreOffice/CoreText substitution bug that drops CJK glyphs during
    headless PDF rendering on this host.
    """
    if not CJK_FONT_PATH.exists():
        raise FileNotFoundError(CJK_FONT_PATH)
    QUOTE_ASSETS.mkdir(parents=True, exist_ok=True)
    digest = hashlib.sha256(text.encode("utf-8")).hexdigest()[:12]
    path = QUOTE_ASSETS / f"{key}-{digest}.png"
    font = ImageFont.truetype(str(CJK_FONT_PATH), font_size_px)
    probe = Image.new("RGBA", (width_px, 100), (255, 255, 255, 0))
    probe_draw = ImageDraw.Draw(probe)
    padding_x = 5
    padding_y = 3
    lines = _wrap_mixed_language_text(
        probe_draw,
        text,
        font,
        width_px - (padding_x * 2),
    )
    sample_box = probe_draw.textbbox((0, 0), "Ag中文", font=font)
    line_height = max(font_size_px + 5, sample_box[3] - sample_box[1] + 5)
    height_px = (padding_y * 2) + (line_height * len(lines))
    image = Image.new("RGBA", (width_px, height_px), (255, 255, 255, 0))
    draw = ImageDraw.Draw(image)
    fill = tuple(bytes.fromhex(INK)) + (255,)
    for line_index, wrapped_line in enumerate(lines):
        draw.text(
            (padding_x, padding_y + (line_index * line_height)),
            wrapped_line,
            font=font,
            fill=fill,
        )
    image.save(path, dpi=(300, 300), optimize=True)
    return path


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text_node, end])
    set_run_font(run, size=8, color=MUTED)


def configure_header_footer(section, first_section=False):
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)
    section.different_first_page_header_footer = first_section
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.text = ""
    set_run_font(p.add_run("PROFAB  /  PROJECT LIFECYCLE CONTROL REPORT"), size=7.5, bold=True, color=BLUE)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    set_run_font(p.add_run("Controlled review copy  •  30 August 2026  |  "), size=8, color=MUTED)
    add_page_number(p)
    if first_section:
        section.first_page_header.paragraphs[0].text = ""
        section.first_page_footer.paragraphs[0].text = ""


def portrait_geometry(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    configure_header_footer(section)


def landscape_geometry(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.48)
    section.left_margin = Inches(0.48)
    section.right_margin = Inches(0.48)
    configure_header_footer(section)


def new_portrait(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    portrait_geometry(section)
    return section


def new_landscape(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    landscape_geometry(section)
    return section


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Title", 30, NAVY, 0, 10),
        ("Heading 1", 21, NAVY, 10, 8),
        ("Heading 2", 14, BLUE, 9, 5),
        ("Heading 3", 10.5, TEAL, 7, 3),
    ]:
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = FONT_LATIN
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    caption.font.size = Pt(8)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(7)


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    keep_with_next(p)
    set_run_font(p.add_run(text.upper()), size=8, bold=True, color=TEAL)
    return p


def add_body(doc, text, bold_lead=None, color=INK, size=9.2, italic=False):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_run_font(p.add_run(bold_lead), size=size, bold=True, color=color)
        set_run_font(p.add_run(text[len(bold_lead):]), size=size, color=color, italic=italic)
    else:
        set_run_font(p.add_run(text), size=size, color=color, italic=italic)
    return p


def add_bullets(doc, items, level=0, color=INK, size=9):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(2.5)
        set_run_font(p.add_run(item), size=size, color=color)


def add_callout(doc, title, body, tone="blue"):
    palette = {
        "blue": (PALE_BLUE, BLUE),
        "teal": (PALE_TEAL, GREEN),
        "amber": (PALE_AMBER, AMBER),
        "red": (PALE_RED, RED),
        "grey": (PALE_GREY, MUTED),
    }
    fill, accent = palette[tone]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run(title), size=10, bold=True, color=accent)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    set_run_font(p.add_run(body), size=8.8, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_metric_row(doc, metrics):
    table = doc.add_table(rows=1, cols=len(metrics))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for idx, (value, label, tone) in enumerate(metrics):
        cell = table.cell(0, idx)
        shade_cell(cell, tone)
        set_cell_margins(cell, top=120, start=100, bottom=115, end=100)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(value), size=19, bold=True, color=NAVY)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(label), size=7.2, bold=True, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_table(doc, headers, rows, widths=None, font_size=7.8, header_fill=NAVY,
              alternate=True, first_col_bold=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_fixed(table)
    header = table.rows[0]
    repeat_header(header)
    prevent_row_split(header)
    for col, text in enumerate(headers):
        cell = header.cells[col]
        shade_cell(cell, header_fill)
        set_cell_margins(cell, top=80, start=90, bottom=80, end=90)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(str(text)), size=font_size, bold=True, color=WHITE)
        if widths:
            cell.width = Inches(widths[col])
    for row_idx, row_data in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        if alternate and row_idx % 2:
            for cell in row.cells:
                shade_cell(cell, PALE_GREY)
        for col, value in enumerate(row_data):
            cell = row.cells[col]
            set_cell_margins(cell, top=65, start=80, bottom=65, end=80)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if widths:
                cell.width = Inches(widths[col])
            if isinstance(value, list):
                set_cell_lines(cell, value, size=font_size)
            else:
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.02
                set_run_font(
                    p.add_run(str(value)),
                    size=font_size,
                    bold=first_col_bold and col == 0,
                    color=INK,
                )
    return table


def add_figure(doc, filename, caption, width_inches, figure_no):
    path = SCREENSHOTS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(cap.add_run(f"Figure {figure_no}. {caption}"), size=8, italic=True, color=MUTED)


def parse_forms():
    records = []
    for line in FORM_SOURCE.read_text(encoding="utf-8").splitlines():
        if not line.strip().startswith("define({"):
            continue

        def value(key):
            match = re.search(rf'{re.escape(key)}:\s*"([^"]*)"', line)
            return match.group(1) if match else ""

        def string_list(key):
            match = re.search(rf'{re.escape(key)}:\s*\[([^\]]*)\]', line)
            return re.findall(r'"([^"]+)"', match.group(1)) if match else []

        included_match = re.search(r'\.\.\.included\("([^"]+)"(?:,\s*"([^"]+)")?\)', line)
        if included_match:
            availability = "Included"
            pages = included_match.group(1)
            version = included_match.group(2) or ""
        elif 'sourceAvailability: "Supplemental"' in line:
            availability = "Supplemental"
            pages = value("sourcePages")
            version = value("sourceVersion")
        else:
            availability = "Index Only"
            pages = ""
            version = ""

        record = {
            "index": value("index"),
            "code": value("code"),
            "title": value("title"),
            "stage": value("stage"),
            "home": value("linkedLayer2NodeId"),
            "touchpoints": string_list("lifecycleTouchpoints"),
            "type": value("type"),
            "profile": value("profile"),
            "description": value("description"),
            "focus": value("focusLabel"),
            "owner": value("responsibleRole"),
            "approvers": string_list("approvalRoles"),
            "applicability": value("defaultApplicability"),
            "signature": "signatureRequired: true" in line,
            "approval": "approvalRequired: true" in line,
            "availability": availability,
            "pages": pages,
            "version": version,
            "source_note": value("sourceNote"),
        }
        records.append(record)
    if len(records) != 52:
        raise RuntimeError(f"Expected 52 controlled forms, parsed {len(records)}")
    return records


L1_ROWS = [
    ("1", "Start", "INITIAL CONTACT", "Project Start"),
    ("2", "Phase", "OPPORTUNITY & QUALIFICATION", "Opportunity & Qualification"),
    ("3", "Primary Gate", "G1 — QUALIFIED & COMMERCIALLY ENGAGED", "G1 — Qualified & Commercially Engaged"),
    ("4", "Phase", "PRE-CONSTRUCTION", "Commercial Pathway + Pre-Construction"),
    ("5", "Primary Gate", "G2 — PROJECT / TECHNICAL COMMITMENT", "G2 — Project / Technical Commitment"),
    ("6", "Phase", "PRODUCTION READINESS", "Production Readiness"),
    ("7", "Primary Gate", "G3 — PRODUCTION AUTHORIZATION", "G3 — Production Authorization"),
    ("8", "Phase", "FACTORY PRODUCTION", "Factory Production"),
    ("9", "Primary Gate", "G4 — FACTORY COMPLETION / RELEASE", "G4 — Factory Completion / Release"),
    ("10", "Phase", "DELIVERY / PROJECT COMPLETION", "Delivery / Project Completion"),
    ("11", "Primary Gate", "G5 — PROJECT COMPLETION / WARRANTY START", "G5 — Project Completion / Warranty Start"),
    ("12", "Phase", "COMMISSIONING & WARRANTY", "Commissioning & Warranty"),
    ("13", "End", "FINAL CLOSE", "Final Close"),
]

L2_ROWS = [
    ("Project Start", "Entry record and project identifier", "L1 Initial Contact"),
    ("Opportunity & Qualification", "Six-step Gate 1 evidence and routing workspace", "L1 Opportunity"),
    ("G1 — Qualified & Commercially Engaged", "Independent release point; no bypass", "L1 G1"),
    ("Commercial Pathway", "Bounded CSA → Class D → PCS, governed LOI modifier", "L1 Pre-Construction"),
    ("Pre-Construction", "Design, site, estimating, scope and responsibilities", "L1 Pre-Construction"),
    ("G2 — Project / Technical Commitment", "Class C basis and project/technical commitment", "L1 G2"),
    ("Production Readiness", "Drawings, PSO, materials, permit, payment, schedule and changes", "L1 Production Readiness"),
    ("G3 — Production Authorization", "Controlled factory release", "L1 G3"),
    ("Factory Production", "Production, QC, inspection and rework", "L1 Factory Production"),
    ("G4 — Factory Completion / Release", "Completion evidence plus transport/site readiness", "L1 G4"),
    ("Delivery / Project Completion", "Site, delivery, set, interfaces, deficiencies", "L1 Delivery"),
    ("G5 — Project Completion / Warranty Start", "Acceptance, deficiency disposition and warranty basis", "L1 G5"),
    ("Commissioning & Warranty", "Startup, handover, claims and expiry tracking", "L1 Commissioning & Warranty"),
    ("Final Close", "Commercial and obligation reconciliation", "L1 Final Close"),
]


REQUIREMENTS = [
    ("§3", "Process before Forms — 先把 end-to-end process 定义清楚，再决定 Forms 放哪、谁签、什么时候签、怎么 version。", "System architecture: L1 defines the whole lifecycle; L2 defines process and Gates; L3 owns forms, signatures, approvals and versions. Evidence: Figs. 1–4."),
    ("§3", "Main Gates + Secondary Controls — Primary Gates 少而清楚；复杂 Technical / Commercial / Payment / Design controls 放到下面。", "L1 contains exactly five primary Gates. Technical, commercial, payment, design and release conditions stay in L2/L3. Evidence: Figs. 1–3."),
    ("§3", "Gate ≠ Document — Gate 是 decision / control point，不是某一份表格的代名词。", "Five Gate nodes are independent L2 release points. Forms supply evidence but no form is used as the Gate itself."),
    ("§3", "Dynamic but Bounded — 不同客户可以走不同 path，但 variations 必须被限制在有限 acceptable patterns，不能 spaghetti。", "Commercial Pathway limits early routing to CSA, Class D, PCS, governed LOI, Hold or No-Go; direct Opportunity → Commercial bypass is removed. Evidence: Fig. 3."),
    ("§3", "Evidence, not Feelings — Sales 的“感觉客户靠谱”不够；要转成 Name / Role / Authority / Budget / Site / Design / Consultant 等可验证事实。", "Gate 1 collects named authority, role, approval path, scale, site, design, budget, funding and team evidence. Missing authority evidence is a hard Hold. Evidence: Fig. 5."),
    ("§3", "Whole Project, not Module Only — 即使 Civil/Foundation/Site/Commissioning 不是 ProFab scope，也必须在 overall process 中被追踪。", "L1/L2 include site, delivery, commissioning, warranty and Final Close; the 18-row Responsibility Matrix tracks civil/foundation/site interfaces."),
    ("§3", "Simple Input, Useful Output — 员工输入要简单；系统输出要能回答 management 真正关心的问题。", "Progressive L1 → L2 → L3 disclosure; status summaries roll from L3 to L2. Gate 1 explains route, Hold reason, risks and next action."),
    ("§3", "Controlled Master — Forms、语言版本、Terms、编号、模板必须受控，避免 Sales / English / French 版本漂移。", "Every catalog form has a locked index/code/title, controlled revision and language. Matrix/source QA exposes discrepancies. Evidence: Figs. 4 and 6."),
    ("§3", "Break the System Before Launch — 主动找 loophole、绕过路径和冲突点，正式上线前先尝试把流程“弄坏”。", "63 rule tests, 9 interaction tests and 7 responsive/accessibility tests exercise hard stops, route abuse, form bypasses and UI regressions."),
    ("§4", "[明确要求] High-Level 必须从 Initial Contact 一直看到 End of Warranty / Final Close；必须是完整 project，而不是只画 module production。High-Level 要简单、visual、让新人 / Sales / Management 都能一眼理解。", "L1 is the exact 13-node Initial Contact → Final Close view with five highlighted primary Gates. Evidence: Figs. 1–2."),
    ("§4", "[明确要求] Primary Gates 要少；详细 Design Approval、Sales Agreement、Shop Drawing Start、Procurement Release、Payment Milestones 等不能全部升级成 L1 Primary Gate。", "Only G1–G5 are primary Gates. PDAF, SA, PSO, payment and procurement controls remain linked L3 records/conditions."),
    ("§5", "Primary Gate 应回答“项目能不能继续 / 谁可以 release / 什么状态发生改变”。", "Each Gate has release conditions, an approval execution item, responsible authority and GO/HOLD status aggregation."),
    ("§5", "Documents 可以跨多个 Gate 重复出现，只要它还 active、会 update 或 later trigger。", "Each form has lifecycle touchpoints; CEC, SOW, RM, changes, incidents and warranty records link across multiple L2 stages."),
    ("§6", "[明确要求] Gate 1 的工作不是简单“打分过线”，而是先定义 acceptable client / owner、收集客观 Evidence、识别 fundamental blocker，并知道这个客户接下来应该走哪条 path。", "Opportunity & Qualification implements six ordered steps: evidence, hard blockers, plan qualification, route, LOI governance, Gate dossier. Score cannot release a hard stop."),
    ("§6", "[明确要求] 通过 Gate 1 后，系统至少应该知道：这个 Client / Owner 是谁、谁能决定、Project 基本是什么、Site / Design / Budget / Funding / Consultants 到什么程度，以及下一步 Route 是什么。", "G1 handoff requires identity, full authority evidence, storeys + GFA, site, design and assigned/executed eligible route; other facts remain visible as actions/risks."),
    ("§6.1", "Number of Storeys + Approx. GFA 是 Class D 最关键的早期规模输入；Module Count 早期可以 optional。", "Storeys and GFA are mandatory positive values for G1/Class D readiness; either missing value blocks. Module count remains optional."),
    ("§6.1", "必须和 Design Maturity 分开；Permit Issued 不代表 modular compatible。", "Design maturity and modular compatibility are separate fields/rules. Permit/IFC without modular review requires Technical Review and blocks G1."),
    ("§6.1", "Grant / financing / equity / government / mixed 等；未完全落实通常是 Risk，不自动 No-Go。", "Funding In Process, Not Secured and Unknown produce visible risk flags and actions; none independently produces No-Go."),
    ("§6.2", "Evidence / Intake：收客观事实。\nRules：Hard / Conditional / Risk。\nEligibility：不是问 Pass/Fail，而是问 Eligible for WHAT。\nRoute：CSA / PCS / governed LOI / Class D / Technical Review / Hold / No-Go 等。\nScore：最后才算，作为 management quality indicator；永远不能覆盖 Hard Rule。", "The rule engine follows this exact order. Eligibility and route are deterministic; score is calculated last and never overrides a Hard rule."),
    ("§6.2", "典型要求：Decision Authority 完全 Unknown 可以 HOLD；No Design 不自动 No-Go；Funding In Process 通常是 Risk；Permit Issued 但 modular incompatible 不能因为“设计成熟”而通过。", "All four examples are encoded and tested: authority Hold; No Design → CSA; funding risk; permit-incompatible → Hold/No-Go."),
    ("§6.3", "[明确要求] 不同客户起点不同，Gate 1 / early pre-con 不能只有一条 linear route；但路径必须 bounded。", "Candidate/Multiple/Unassigned/No Site routes to bounded CSA; plan routes to Class D; recorded Class D routes to PCS; governed LOI cannot shortcut the sequence."),
    ("§6.3", "可能使用 optional governed LOI，但不能借此绕过后续受控 CSA / PCS / project path。", "LOI is unavailable to Standard relationships unless policy enables it; it requires Class D evidence, scope/time caps, conversion trigger, review date and executive sign-off."),
    ("§7", "Class D 不是 quote；是 Sales-level early qualification / rough budget tool。", "Gate 1 records route readiness only and does not require or invent a Class D amount. The controlled CEC form records the later estimate basis and disclaimer evidence."),
    ("§7", "如果连 Stories / GFA 都不知道，不应硬做一个假 Class D；应回到 Pre-Qualification / Consultation 补足项目规模。", "Missing or non-positive storeys/GFA blocks qualification. No synthetic estimate or default $/sq.ft. amount is generated."),
    ("§8", "[明确 / strong working direction] 前面的客户可以走不同路线，但一旦达到 Class C，应该尽量成为“同一种可估价项目状态”；之前属于哪种 Client Type 就不再重要。", "Commercial Pathway converges early variants into the Pre-Construction/CEC stream; G2 accepts a common Class C-ready technical/scope basis."),
    ("§9", "流程必须可以 re-evaluate / loop / hold / resume；但 coordinator 不应必须靠“某个作者本人”才能理解路径。", "Opportunity can be re-evaluated at any step; status/rules are derived from saved evidence. L3 recurring/triggered controls preserve loops and exceptions without spaghetti."),
    ("§10", "Signed PSO + applicable funds 可 unlock Purchasing 的现有业务逻辑；Product decisions 必须有受控 release。", "PSO is a required signed/approved L3 control linked to Production Readiness and G3; final funds/approver policy remains a documented TBD."),
    ("§11", "Production 前必须有真正的 GO / HOLD control；Production 不应自己去多个文件里猜“到底 ready 了没有”。", "Production Readiness aggregates drawings, PSO, materials, permit, payment, MPS and changes; G3 is a separate approval release."),
    ("§11", "Factory work complete 不代表自动可以 shipment；release 还涉及 module condition、payment、logistics、site readiness / foundation / access / crane 等 whole-project dependency。", "G4 separates Factory Completion from Release. MSO, QC, inspection, transport and site-readiness evidence remain prerequisites."),
    ("§12", "Warranty 不应该只是“文件最后一张”；系统要能追踪 warranty start basis、open obligations、claim/support、expiry，直到 Final Close。", "G5 records warranty start; Commissioning & Warranty tracks WMA/support/notices; Final Close reconciles open obligations and acceptance."),
    ("§13.2", "同一 document 可以在多个 stage 出现。\n使用 REQUIRED / OPTIONAL / SUPPORTING / TRIGGERED 等标签，而不是强迫所有文件线性发生。", "All 52 forms carry one home node plus one-or-more lifecycle touchpoints and one of six applicability classes; Conditional/Triggered records require an explicit decision."),
    ("§14", "RM\nHigh-level core-party responsibility allocation，通常 GC / ProFab / Client；不要塞所有 subcontractors。\nSOW\n更详细的 scope delineation：included/excluded、连接边界、谁 supply/install。", "The L3 Responsibility Matrix stays at 18 high-level delivery areas/4 core parties; SOW is a separate detailed controlled form for exact boundaries."),
    ("§19", "界面/输入必须 simple；先高层 visual，再 progressive breakdown。", "L1 is intentionally concise; L2 adds process/gates; L3 exposes forms, matrices, fields, approvals and evidence. Responsive behavior was verified at five widths."),
    ("§21", "然后把现有 documents 映射到真实 process，找缺失 / 重复 / recurring / triggered files。\n再标准化 / 创建 Forms，处理 version / bilingual / signature / document control。\n用 past projects 回放测试 process，主动找 loopholes。", "Complete 52-record L3 catalog, source QA, revision/language/signature/approval controls, plus automated replay scenarios and full UI tests."),
]


SCENARIOS = [
    ("Q01", "Valid confirmed authority + Plan + scale", "Class D route; G1 may pass without a pre-G1 estimate amount", "PASS — Class D; G1 passes"),
    ("Q02", "No Design", "Route to paid CSA; never automatic No-Go", "PASS — CSA; G1 can pass"),
    ("Q03", "Decision authority Unknown", "Hard Hold regardless of score", "PASS — HOLD"),
    ("Q04", "Decision authority Partially Confirmed", "Hard Hold; cannot be treated as final", "PASS — HOLD"),
    ("Q05", "Final authority explicitly No", "Hard Hold", "PASS — HOLD"),
    ("Q06", "Authority marked Confirmed; name/role blank", "Objective-evidence Hold", "PASS — HOLD"),
    ("Q07", "Decision parties Yes; approval path blank", "Approval-path Hold", "PASS — HOLD"),
    ("Q08", "Required decision parties No/Unknown", "Hard Hold", "PASS — HOLD"),
    ("Q09", "Storeys and GFA both missing", "Block; no fake Class D", "PASS — BLOCKED"),
    ("Q10", "Storeys missing; GFA present", "Block incomplete scale", "PASS — BLOCKED"),
    ("Q11", "GFA missing; storeys present", "Block incomplete scale", "PASS — BLOCKED"),
    ("Q12", "Candidate Site", "Bounded CSA route + site risk", "PASS — CSA"),
    ("Q13", "Multiple Sites", "Bounded CSA route + site risk", "PASS — CSA"),
    ("Q14", "Municipal land not assigned", "Bounded CSA route + site risk", "PASS — CSA"),
    ("Q15", "No Site", "Bounded CSA/site-feasibility route", "PASS — CSA"),
    ("Q16", "Fatal site/logistics issue, not resolvable", "No-Go; archive route", "PASS — NO-GO"),
    ("Q17", "Fatal issue potentially resolvable", "Hold pending approved resolution", "PASS — HOLD"),
    ("Q18", "Permit Issued + Not Compatible + no corrective path", "No-Go; maturity cannot override compatibility", "PASS — NO-GO"),
    ("Q19", "Permit Issued + Not Reviewed", "Technical Review required; G1 blocked", "PASS — TECH REVIEW"),
    ("Q20", "IFC/Construction Ready + Not Reviewed", "Technical Review required; G1 blocked", "PASS — TECH REVIEW"),
    ("Q21", "Partially Compatible", "Technical Review required", "PASS — TECH REVIEW"),
    ("Q22", "Major Rework Likely", "Technical Review required", "PASS — TECH REVIEW"),
    ("Q23", "Not Compatible + corrective path Unknown", "Hold for technical decision", "PASS — HOLD"),
    ("Q24", "Funding In Process", "Risk only; not automatic No-Go", "PASS — risk visible"),
    ("Q25", "Funding Not Secured", "Distinct risk; no false ‘in process’ label", "PASS — risk visible"),
    ("Q26", "Funding Unknown", "Risk and action", "PASS — risk visible"),
    ("Q27", "Client budget missing", "Action; no synthetic price; not No-Go", "PASS — ACTION REQUIRED"),
    ("Q28", "Budget amount present; basis missing", "Commercial action; route preserved", "PASS — action visible"),
    ("Q29", "Class D shows major budget gap", "Commercial review action", "PASS — Major Gap"),
    ("Q30", "Aggressive/Unrealistic/Review timeline", "Visible risk; no silent pass", "PASS — three variants"),
    ("Q31", "Existing controlled Class D result", "Advance primary route to PCS", "PASS — PCS"),
    ("Q32", "Stale Class D amount while availability = No", "Ignore stale amount; cannot bypass to PCS", "PASS — Class D remains"),
    ("Q33", "No Design plus stale Class D amount", "CSA remains primary route", "PASS — CSA"),
    ("Q34", "Standard client selects LOI; policy disabled", "Route unavailable; G1 blocked", "PASS — blocked"),
    ("Q35", "Returning client selects LOI before Class D", "Cannot shortcut Class D", "PASS — blocked"),
    ("Q36", "Returning client + Class D + bounded LOI", "LOI available only with governance", "PASS — eligible"),
    ("Q37", "LOI missing caps/trigger/review/sign-off", "G1 blocked", "PASS — blocked"),
    ("Q38", "Assigned route does not match eligible route", "G1 blocked", "PASS — blocked"),
    ("Q39", "Commercial instrument not Executed", "G1 blocked", "PASS — blocked"),
    ("Q40", "High score + fatal hard rule", "Hard rule wins", "PASS — NO-GO"),
    ("F01", "Form status Complete; required fields blank", "Effective status remains incomplete", "PASS — blocked"),
    ("F02", "All fields complete; approval pending", "Cannot complete", "PASS — incomplete"),
    ("F03", "Agreement executed; signature pending", "Cannot complete", "PASS — incomplete"),
    ("F04", "Conditional/Triggered applicability Pending", "Gates until determination", "PASS — incomplete"),
    ("F05", "N/A selected; reason blank/whitespace", "N/A cannot bypass", "PASS — incomplete"),
    ("F06", "N/A selected with objective reason", "Non-applicability accepted", "PASS — complete"),
    ("F07", "Approval or signature Rejected", "Hard blocked state", "PASS — blocked"),
    ("F08", "Attempt to delete a controlled form", "Deletion control unavailable", "PASS — protected"),
]


TBD_ITEMS = [
    ("5 个 Primary Gate 的最终名称、边界、Gate Authority / GO-HOLD-RELEASE 权限。", "Working G1–G5 names and approval rows are visible; final authority remains management-controlled."),
    ("G2 的准确 trigger。", "G2 working conditions are present; exact contractual trigger is not declared final."),
    ("Sales Agreement 的准确签署位置和与 supporting documents 的 hierarchy。", "SA is mapped across Commercial/G2/Readiness as Conditional; legal hierarchy remains open."),
    ("Shop Drawing Start 的正式 authorization / payment trigger。", "Represented as an L2 hard-control topic; no invented payment threshold."),
    ("Permit Submission / Permit Issued / IFC 中哪个状态与 shop drawing / technical commitment 正式关联。", "All maturities are tracked separately from compatibility; final trigger remains open."),
    ("Class B / A 与 Issue for Permit / Issue for Construction 的正式对应关系。", "CEC supports D/C/B/A evidence; no permanent B/A policy mapping is fabricated."),
    ("Gate 1 最终 Mandatory / Optional / Conditional criteria 与 Owner Types。", "Current minimum objective controls are implemented; configurable policy can evolve after sign-off."),
    ("No confirmed budget 但愿意付 feasibility / consultation 的客户能否过 Gate 1，以及限制。", "Current engine permits paid-route qualification with a visible budget action; final policy remains open."),
    ("Class C 的精确最低 entry requirements。", "G2 contains a working Class C-ready basis; exact minimum remains a management decision."),
    ("Class D 的最终 assumptions / budget tolerance / ownership。", "No formula/default is invented; tolerance is a configurable rule and CEC records assumptions."),
    ("PCS future fee basis（JF 已质疑 module-based fee，讨论过 sq.ft. direction，但未 final）。", "PCS fee basis is document-controlled, not hard-coded."),
    ("特殊客户 / Buttcon-type 低 preconstruction fee arrangement 的政策边界；现有个别金额不应自动成为公司政策。", "No customer-specific amount or policy is encoded."),
    ("CSA / PCS / LOI 的 entry logic、signing authority、LOI governance。", "Working bounded logic and signatures exist; business-rule/signing policy stays configurable."),
    ("Production Release / Purchasing Release 的绝对 prerequisite 和最终 approver。", "G3/PSO/MPS controls are present; absolute prerequisite/approver awaits sign-off."),
    ("Payment milestones 与各 release 的正式 hard-gate mapping。", "Payment remains a release condition placeholder; no amount or milestone is fabricated."),
    ("Warranty start / commissioning / handover 的正式 contractual trigger，以及不同 warranty 是否不同起点。", "G5/WMA/closeout track the basis and dates; contractual trigger remains project/policy controlled."),
    ("Customer / Project / Job Number / Development / Building hierarchy 与 Job Number trigger。", "Current project identifier supports the workflow; final enterprise hierarchy remains separate work."),
    ("Time-tracking alert threshold、recipient、action、client notification logic。", "Lifecycle status supports future reporting; threshold/notification policy is not invented in this release."),
    ("最终 IT deployment / corporate integration。", "Proof of concept remains local/web-ready; enterprise integration is intentionally outside this deliverable."),
]


BUG_ROWS = [
    ("L1 structure drift", "G1 was not a distinct L1/L2 control and the commercial path could appear as a peer L1 phase.", "Restored exact 13-node L1; Commercial Pathway lives under Pre-Construction; G1 is a distinct Gate."),
    ("Route bypass", "Opportunity could connect directly to Commercial Pathway.", "Canonical edge is Opportunity → G1 → Commercial Pathway; migration removes the bypass."),
    ("Authority-state mismatch", "Partially Confirmed could be visually blocked but pass the underlying Gate decision.", "Partially Confirmed is now a deterministic Hard Hold in both UI and rule engine."),
    ("‘Confirmed’ without evidence", "No name/role or approval path was required.", "Final Authority Name, Role/Title and Approval Path are mandatory hard-control evidence."),
    ("Class D misuse", "A price could be treated as a pre-G1 prerequisite or fabricated from defaults.", "G1 only assigns route; storeys + GFA are mandatory; no amount/formula is invented."),
    ("Maturity shortcut", "Permit/IFC maturity could be mistaken for modular compatibility.", "Separate fields and Technical Review/No-Go rules block incompatible or unreviewed design."),
    ("Controlled-form gaps", "The previous L3 model did not expose every indexed/source form consistently.", "52 unique catalog forms are seeded once and open from their linked L2 nodes."),
    ("Form completion bypass", "Status could be set Complete while required fields, signature or approval were missing.", "Effective progress derives from fields + status + signature + approval + task state."),
    ("N/A bypass", "Conditional/triggered records could disappear without evidence.", "Pending applicability gates; N/A requires an objective reason."),
    ("Controlled form deletion", "A catalog record could be removed and silently reduce the control set.", "Catalog forms and Gate controls cannot be deleted; Clear L3 removes custom items only."),
    ("Uncontrolled collaboration", "The page automatically connected to a public relay.", "Collaboration is Off/Local unless a user explicitly creates or opens a share link."),
    ("Auth-bypass request error", "Local bypass mode still called a missing /api/auth/me endpoint.", "Cloud controls skip that request in bypass mode; no failed response remains."),
    ("Missing favicon", "Browser requested a non-existent icon.", "A controlled app icon now eliminates the 404."),
    ("Unlabelled controls", "11 Gate 1 secondary inputs had visual labels but no programmatic accessible name.", "All 11 were associated with labels/ARIA; deep L3 and matrix scans now pass."),
    ("Responsive density", "Dense L2/L3 layouts risked document-level horizontal overflow.", "L1/L2/L3 verified at 390, 768, 1024, 1440 and 1920 px without document overflow."),
    ("Source naming drift", "SER/MALR/ODN/PCR names and DC/CD/SOW/PSO/WMA numbering could conflict.", "Controlled titles corrected; source discrepancies remain visible as QA notes."),
]


def build_report():
    forms = parse_forms()
    counts = Counter(item["availability"] for item in forms)
    applicability_counts = Counter(item["applicability"] for item in forms)

    doc = Document()
    setup_styles(doc)
    first = doc.sections[0]
    portrait_geometry(first)
    first.different_first_page_header_footer = True
    first.first_page_header.paragraphs[0].text = ""
    first.first_page_footer.paragraphs[0].text = ""
    doc.core_properties.title = "PROFAB Project Lifecycle — JF Requirements Traceability & Verification"
    doc.core_properties.subject = "L1/L2/L3 lifecycle, controlled forms, matrices, testing and QA"
    doc.core_properties.author = "PROFAB workflow implementation review"
    doc.core_properties.keywords = "PROFAB, JF, workflow, lifecycle, Gate 1, L3 forms, traceability"

    # Cover
    cover = doc.add_table(rows=1, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover.autofit = False
    row = cover.rows[0]
    row.height = Inches(8.82)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = row.cells[0]
    shade_cell(cell, NAVY)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell, top=400, start=450, bottom=400, end=450)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    set_run_font(p.add_run("PROFAB"), size=17, bold=True, color=TEAL)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    set_run_font(p.add_run("PROJECT LIFECYCLE"), size=31, bold=True, color=WHITE)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    set_run_font(p.add_run("JF Requirements Traceability & Verification"), size=18, bold=True, color="CDE4F5")
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    set_run_font(p.add_run("L1 clarity  •  L2 process control  •  L3 forms & matrices"), size=10, color="AFC6D9")
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.add_run("CONTROLLED REVIEW COPY"), size=8, bold=True, color=TEAL)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run("30 August 2026  |  Version 1.0"), size=9, color="D5E3EE")
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Prepared for ProFab / JF management review"), size=9, color="D5E3EE")
    doc.add_page_break()

    # Executive assurance
    add_kicker(doc, "01 / Executive assurance")
    doc.add_heading("Lifecycle implemented and verified", level=1)
    add_callout(
        doc,
        "Outcome",
        "The application now implements the source-definable JF lifecycle as a connected three-layer operating model: a simple 13-node L1, a 14-node L2 with a distinct G1 and bounded commercial route, and an L3 containing all 52 controlled form records, two matrices, eight Gate/final controls, and enforceable completion logic.",
        "teal",
    )
    add_metric_row(doc, [
        ("13", "L1 lifecycle nodes", PALE_BLUE),
        ("14", "L2 process nodes", PALE_TEAL),
        ("52", "L3 controlled forms", PALE_AMBER),
        ("79", "functional & UI checks", "F0EEFA"),
    ])
    doc.add_heading("What is complete", level=2)
    add_bullets(doc, [
        "Exact Initial Contact → Final Close high-level lifecycle with five primary Gates.",
        "Evidence → Rules → Eligibility → Route → Score qualification logic; no score can override a hard rule.",
        "All 52 indexed/supplemental documents represented as operational L3 records, each with a 10-field controlled schema.",
        "Approval Matrix: 52 forms + 5 Gates = 57 rows across six approval-role columns.",
        "Responsibility Matrix: 18 whole-project delivery areas across four core-party columns.",
        "Completion, signature, approval, task, applicability and N/A controls tested for bypass resistance.",
        "Production build, console/network health, responsive layouts and accessible control naming verified.",
    ])
    doc.add_heading("Important qualification", level=2)
    add_callout(
        doc,
        "No policy was invented",
        "The provided Master Requirements explicitly leaves several legal, payment, signing and release decisions as TBD. The system exposes working control points but does not silently convert those open decisions into final company policy. Section 9 lists every retained TBD.",
        "amber",
    )
    doc.add_heading("Report map", level=2)
    contents = [
        ("1", "Executive assurance"), ("2", "Evidence basis & source integrity"),
        ("3", "Three-layer lifecycle architecture"), ("4", "JF requirement traceability"),
        ("5", "Complete L3 form & matrix register"), ("6", "Gate 1 enforcement design"),
        ("7", "Scenario and regression verification"), ("8", "Bug & UI review"),
        ("9", "Retained management decisions (TBD)"), ("10", "Acceptance conclusion"),
    ]
    add_table(doc, ["Section", "Contents"], contents, widths=[0.8, 5.9], font_size=8.5, header_fill=BLUE)

    # Source evidence
    doc.add_page_break()
    add_kicker(doc, "02 / Evidence basis")
    doc.add_heading("Source integrity and interpretation rules", level=1)
    add_callout(
        doc,
        "Quotation basis",
        "The left-hand wording in the traceability table is quoted verbatim from JF_Master_Requirements_All_Conversations_2026-08-29.docx. That file describes itself as a compiled Master Requirements summary, not a word-for-word conversation transcript. Accordingly, this report labels the text as source wording from the Master summary—not as a direct transcript quote.",
        "blue",
    )
    sources = [
        ("JF_Master_Requirements_All_Conversations_2026-08-29.docx", "Requirements, working directions and explicit TBDs", "12 rendered pages; internally coherent summary"),
        ("Combined Forms.pdf", "Controlled index, included form bodies, revisions and numbering", "193 pages; forms usable through page 178; pages 179–193 corrupted binary render"),
        ("Current PROFAB Project Lifecycle application", "Existing architecture, implementation target and verification surface", "Production build and browser flows tested locally"),
    ]
    add_table(doc, ["Evidence source", "Use in this review", "Observed condition"], sources, widths=[2.5, 2.6, 1.7], font_size=7.8, header_fill=BLUE)
    doc.add_heading("Controlled-form reconciliation", level=2)
    add_metric_row(doc, [
        ("52", "digital L3 records", PALE_BLUE),
        (str(counts["Included"]), "included templates", PALE_TEAL),
        (str(counts["Index Only"]), "index-only templates", PALE_AMBER),
        (str(counts["Supplemental"]), "supplemental template", "F0EEFA"),
    ])
    add_bullets(doc, [
        "The nine-page PDF register contains 51 indexed entries.",
        "SMO 2.13.1 exists on PDF pages 114–116 but is omitted from that index; it is preserved as the 52nd supplemental controlled record.",
        "The 21 Included records use the source form pages/revisions as evidence; 30 Index Only records receive a controlled operational digital schema without pretending that a missing PDF body was supplied.",
        "PDF pages 179–193 render embedded XLSX/DOCX binary data and are not treated as forms.",
        "DC/CD and SOW/PSO/WMA numbering conflicts are exposed to users rather than silently normalized away.",
    ])
    add_figure(doc, "05-l3-form-register-source-qa.png", "L3 source-package QA visible inside the live Approval Matrix.", 6.7, 1)

    # Architecture
    doc.add_page_break()
    add_kicker(doc, "03 / Three-layer architecture")
    doc.add_heading("L1 is simple; L2 explains; L3 executes", level=1)
    add_callout(doc, "Layer contract", "L1 answers ‘where is the project?’ L2 answers ‘what process and Gate comes next?’ L3 answers ‘what evidence, form, matrix, signature, approval or task makes that status true?’", "teal")
    doc.add_heading("L1 — exact high-level lifecycle", level=2)
    add_figure(doc, "01b-l1-13-step-overview.png", "The complete 13-node L1 lifecycle and five primary Gates.", 6.8, 2)
    add_table(doc, ["#", "Type", "L1 node", "Linked L2 node(s)"], L1_ROWS, widths=[0.35, 0.9, 3.15, 2.25], font_size=7.4, header_fill=NAVY)
    add_figure(doc, "01-l1-project-lifecycle.png", "L1 phase cards show explicit links into their L2 process nodes.", 6.8, 3)

    new_landscape(doc)
    add_kicker(doc, "03 / Three-layer architecture")
    doc.add_heading("L2 — bounded process with independent Gates", level=1)
    add_table(doc, ["L2 node", "Purpose", "L1 parent"], L2_ROWS, widths=[2.8, 4.8, 2.2], font_size=7.6, header_fill=BLUE)
    add_figure(doc, "02-l2-bounded-process.png", "Opportunity cannot bypass G1; the controlled commercial pathway begins only after release.", 9.2, 4)

    new_portrait(doc)
    add_kicker(doc, "03 / Three-layer architecture")
    doc.add_heading("L3 — controlled evidence, not a document dump", level=1)
    add_body(doc, "Every catalog form opens inside the L3 execution layer with a controlled identity, revision, language, applicability decision, lifecycle links, source reference and ten operational fields. Form values are not enough on their own: required signature, approval, task and status controls must also pass.")
    state_rows = [
        ("Applicability", "Required is always applicable; Conditional/Triggered begin Pending; Optional/Supporting do not gate by default."),
        ("Required data", "All required form fields must contain objective values before Complete/Passed becomes selectable."),
        ("N/A", "Not Applicable requires an objective reason; blank or whitespace evidence remains incomplete."),
        ("Signature", "Signed is mandatory where configured; Partially Signed is not complete; Rejected is blocked."),
        ("Approval", "Approved is mandatory where configured; Rejected is blocked."),
        ("Protection", "Catalog forms and Gate controls cannot be deleted; Clear L3 only removes custom items."),
        ("Roll-up", "Only effective L3 progress counts in the L2 status summary."),
    ]
    add_table(doc, ["Control", "Effective behavior"], state_rows, widths=[1.25, 5.5], font_size=8.1, header_fill=TEAL)
    add_figure(doc, "03-l3-controlled-cec-form.png", "A live CEC form shows revision, language, lifecycle links, required fields and blocking status.", 6.75, 5)

    # Traceability table
    new_landscape(doc)
    add_kicker(doc, "04 / Requirement traceability")
    doc.add_heading("JF source wording → implemented node/control", level=1)
    add_body(doc, "The left column preserves the Master Requirements summary wording. The right column identifies the corresponding live node/control and its verification evidence.", size=8.5)
    trace_rows = []
    for requirement_index, (section_ref, quote, implementation) in enumerate(REQUIREMENTS, start=1):
        trace_rows.append((
            [
                {"text": f"{section_ref} · Master summary", "bold": True, "color": BLUE},
                {
                    "image_path": build_source_quote_image(quote, f"requirement-{requirement_index:02d}"),
                    "width_inches": 4.62,
                    "alt_text": quote,
                },
            ],
            [
                {"text": "Implemented", "bold": True, "color": GREEN},
                {"text": implementation, "color": INK},
            ],
        ))
    add_table(doc, ["JF source wording (verbatim from Master Requirements summary)", "Implemented node(s), control and evidence"], trace_rows, widths=[4.9, 5.0], font_size=7.25, header_fill=NAVY, alternate=True)

    # Forms
    new_portrait(doc)
    add_kicker(doc, "05 / L3 controlled register")
    doc.add_heading("Every form has an operational schema", level=1)
    add_body(doc, "The digital form is intentionally more than a file link. Each of the 52 records has six common evidence fields plus four fields selected by its operational profile. If an Optional or Supporting form is used, its required internal fields still apply; its optionality only prevents it from becoming an automatic Gate prerequisite.")
    common_rows = [
        ("All profiles", "Project name; Project/file number; Prepared by; Record date; form-specific primary record; Evidence/attachment references"),
        ("Presentation", "Intended audience; Purpose/call to action; Content revision; Release status"),
        ("Agreement", "Counterparty; Scope/commercial basis; Effective date; Execution state"),
        ("Meeting", "Meeting date; Attendees/organizations; Decisions; Open actions/owners/dates"),
        ("Analysis", "Analysis/estimate basis; Assumptions/exclusions; Findings; Recommendation/next action"),
        ("Approval", "Approval subject/revision; Authorized approver; Approval date; Decision"),
        ("Register", "Record/issue reference; Action owner; Target/response date; Record status"),
        ("Report", "Module/area/subject; Inspection/event date; Findings/evidence; Disposition"),
        ("Instructions", "Applicable revision; Prerequisites; Controlled sequence; Issued to/acknowledged by"),
        ("Schedule", "Baseline date; Period/milestone range; Critical constraints; Schedule status"),
        ("Service", "Service request; Owner; Response date; Service status"),
        ("Package", "Contents/index; Revision/change basis; Transmittal; Acceptance status"),
        ("Notice", "Recipient/authority; Event date; Impact; Response deadline"),
    ]
    add_table(doc, ["Schema", "Controlled fields"], common_rows, widths=[1.35, 5.4], font_size=7.8, header_fill=TEAL)
    add_body(doc, f"Applicability distribution: " + ", ".join(f"{key} {value}" for key, value in sorted(applicability_counts.items())) + ".", size=8.2, italic=True, color=MUTED)

    new_landscape(doc)
    add_kicker(doc, "05 / L3 controlled register")
    doc.add_heading("Complete 52-form register", level=1)
    form_rows = []
    for form in forms:
        source = form["availability"]
        if form["pages"]:
            source += f" pp. {form['pages']}"
        if form["version"]:
            source += f"\n{form['version']}"
        if form["source_note"]:
            source += f"\nQA: {form['source_note']}"
        controls = f"{form['applicability']} · {form['type']}\n10 fields"
        if form["signature"]:
            controls += " · signature"
        if form["approval"]:
            controls += " · approval"
        touchpoints = " → ".join(SHORT_NODE_LABELS.get(item, item) for item in form["touchpoints"])
        form_rows.append((
            f"{form['index']}\n{form['code']}",
            [
                {"text": form["title"], "bold": True, "color": INK},
                {"text": f"Focus: {form['focus']}", "italic": True, "color": MUTED},
            ],
            NODE_LABELS.get(form["home"], form["home"]),
            touchpoints,
            controls,
            [
                {"text": source, "color": INK},
                {"text": f"Owner: {form['owner']}", "color": MUTED},
                {"text": f"Approvers: {', '.join(form['approvers'])}", "color": MUTED},
            ],
        ))
    add_table(
        doc,
        ["Index / code", "Controlled title & operational focus", "L2 home", "Lifecycle touchpoints", "Applicability / controls", "Source / owner / approvers"],
        form_rows,
        widths=[0.7, 2.35, 1.55, 2.0, 1.4, 2.0],
        font_size=6.2,
        header_fill=NAVY,
        alternate=True,
    )

    new_portrait(doc)
    add_kicker(doc, "05 / L3 matrices & release controls")
    doc.add_heading("Matrices connect documents to accountable roles", level=1)
    matrix_rows = [
        ("Approval Matrix — Forms & Actions", "57 rows", "Sales; Technical; Project Mgmt; Factory/Site; Management; Client/Consultants", "52 controlled forms + G1–G5"),
        ("RM — Responsibility Matrix", "18 rows", "GC/Builder; ProFab/Guildcrest; Client/Owner; Consultants", "Whole-project responsibilities from authority/design through warranty/final close"),
    ]
    add_table(doc, ["Matrix", "Size", "Columns", "Purpose"], matrix_rows, widths=[1.7, 0.65, 2.5, 1.95], font_size=7.4, header_fill=VIOLET)
    doc.add_heading("Eight non-form L3 controls", level=2)
    add_bullets(doc, [
        "Gate 1 objective evidence dossier",
        "Gate 1 commercial engagement confirmation",
        "G1 qualified & commercially engaged approval",
        "G2 project / technical commitment approval",
        "G3 production authorization",
        "G4 factory completion / release approval",
        "G5 project completion / warranty start approval",
        "Final close commercial and obligation reconciliation",
    ], size=8.4)

    # Gate 1
    doc.add_page_break()
    add_kicker(doc, "06 / Gate 1 client enforcement")
    doc.add_heading("Evidence—not score—stops non-compliant clients", level=1)
    gate_rows = [
        ("1. Objective Intake", "Client, relationship, authority status, storeys, GFA, site and design basis are recorded."),
        ("2. Hard Blockers", "Authority, full decision group, named role evidence, approval path, scale, fatal site and modular incompatibility are enforced."),
        ("3. Plan Qualification", "No design → CSA; Plan + scale → Class D; no amount is invented."),
        ("4. Commercial Routing", "Assigned route must be one of the currently eligible bounded routes."),
        ("5. LOI Governance", "When used: eligible relationship/policy, named executive, bounded scope/days/hours, review date and paid conversion trigger."),
        ("6. G1 Handoff", "All prior checks plus an Executed commercial instrument; dossier records rules, route, score, risks and open actions."),
    ]
    add_table(doc, ["Gate 1 step", "Release logic"], gate_rows, widths=[1.45, 5.3], font_size=8.0, header_fill=RED)
    add_body(doc, "Hard-rule precedence: The score is a management quality indicator only. A 100/100 opportunity remains Hold, Blocked or No-Go whenever a hard rule is active.", bold_lead="Hard-rule precedence:", size=8.3, color=RED)
    add_figure(doc, "04-g1-objective-evidence-hard-block.png", "Confirmed authority cannot pass without a named role and documented approval path; Eligibility Cleared remains disabled.", 6.2, 6)

    # Scenario verification
    new_landscape(doc)
    add_kicker(doc, "07 / Verification")
    doc.add_heading("Representative enforcement scenario matrix", level=1)
    add_body(doc, "The table below lists 48 representative qualification and form-bypass paths. The automated rule suite contains 63 checks; browser interaction and UI suites add 16 further checks.", size=8.3)
    scenario_table = add_table(
        doc,
        ["ID", "Scenario", "Expected enforcement", "Verified result"],
        SCENARIOS,
        widths=[0.55, 3.1, 3.8, 2.35],
        font_size=6.9,
        header_fill=GREEN,
        alternate=True,
        first_col_bold=True,
    )
    for row in scenario_table.rows[1:]:
        shade_cell(row.cells[3], PALE_TEAL)
        for paragraph in row.cells[3].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = rgb(GREEN)
                run.bold = True

    new_portrait(doc)
    add_kicker(doc, "07 / Verification")
    doc.add_heading("Automated and manual QA evidence", level=1)
    qa_rows = [
        ("Deterministic rules", "63/63 passed", "Qualification hard rules, routing, LOI, score precedence, L1/L2/L3 structure, 52-form catalog, completion/N/A/signature/approval loopholes"),
        ("Browser interaction", "9/9 passed", "App load, exact lifecycle, all 52 forms opening, controlled-form regression, N/A, document metadata, client hard block, matrix, mobile"),
        ("Responsive & accessibility", "7/7 passed", "L1/L2/L3 at five widths; visible inputs/buttons named; duplicate IDs absent on Gate, form and matrix screens"),
        ("Report screenshots", "5/5 passed", "Deterministic capture of L1, L2, L3, hard blocker and source QA"),
        ("Static quality", "Passed", "TypeScript, ESLint, git diff check"),
        ("Production build", "Passed", "Next.js 16.3 optimized static build"),
        ("Console/network", "Passed", "No browser console errors and no HTTP 4xx/5xx on clean load"),
    ]
    add_table(doc, ["Verification layer", "Result", "Coverage"], qa_rows, widths=[1.45, 1.05, 4.25], font_size=7.8, header_fill=GREEN)
    add_callout(doc, "Dependency note", "npm audit reports two Moderate entries that are the same transitive chain: exceljs 4.4.0 → uuid 8.3.2. ExcelJS uses uuid v4 without the advisory's v3/v5/v6 buffer argument, so the current export path is not the affected call pattern. No forced downgrade was applied because npm's suggested fix would move ExcelJS to 3.4.0 and risks breaking exports. There are zero High or Critical findings.", "amber")

    # Bug and UI review
    new_landscape(doc)
    add_kicker(doc, "08 / Bug & UI review")
    doc.add_heading("Issues found and corrected", level=1)
    add_table(doc, ["Finding", "Failure risk", "Correction in current system"], BUG_ROWS, widths=[1.8, 3.55, 4.65], font_size=7.1, header_fill=BLUE, alternate=True)

    new_portrait(doc)
    add_kicker(doc, "08 / UI acceptance")
    doc.add_heading("Current UI acceptance result", level=1)
    add_metric_row(doc, [
        ("5", "verified viewport widths", PALE_BLUE),
        ("0", "document overflows", PALE_TEAL),
        ("0", "unnamed visible controls", PALE_AMBER),
        ("0", "clean-load request errors", "F0EEFA"),
    ])
    add_bullets(doc, [
        "L1 remains visually simple and includes a compact 13-step overview.",
        "L2 makes G1 and the bounded Commercial Pathway visually explicit.",
        "L3 exposes all documents, matrices, source notes and operational fields without changing L2 into a document dump.",
        "The mobile layout uses internal scrolling where required but never causes document-level horizontal overflow.",
        "Collaboration state is clearly labelled Off/Local until explicitly activated.",
        "Controlled forms show immediate blocker text when values regress after completion.",
    ])

    # TBD
    new_landscape(doc)
    add_kicker(doc, "09 / Management decisions retained")
    doc.add_heading("Explicit TBDs were preserved—not guessed", level=1)
    add_body(doc, "These items are copied from the Master Requirements TBD list. They require JF/Management/legal/commercial approval before the working control becomes final policy.", size=8.4)
    tbd_rows = [
        ([{
            "image_path": build_source_quote_image(source, f"tbd-{item_index:02d}"),
            "width_inches": 4.62,
            "alt_text": source,
        }], treatment)
        for item_index, (source, treatment) in enumerate(TBD_ITEMS, start=1)
    ]
    add_table(doc, ["JF / Management TBD (source wording)", "Current system treatment"], tbd_rows, widths=[4.9, 5.0], font_size=7.2, header_fill=AMBER, alternate=True)

    # Conclusion
    new_portrait(doc)
    add_kicker(doc, "10 / Acceptance conclusion")
    doc.add_heading("Ready for controlled review and project replay", level=1)
    add_callout(
        doc,
        "Acceptance position",
        "All source-definable lifecycle, Gate 1, form-register, matrix and completion-control requirements in scope are implemented and verified. The remaining limitations are not hidden software defects: they are damaged/missing source artifacts, explicitly open management policies, and enterprise modules intentionally outside this lifecycle release.",
        "teal",
    )
    doc.add_heading("Recommended next controlled actions", level=2)
    add_bullets(doc, [
        "JF/Management reviews and signs the 19 retained TBD decisions in Section 9.",
        "Document Control replaces the 30 index-only records with approved bilingual source masters as those templates are formally issued.",
        "The source owner repairs/re-exports Combined Forms.pdf pages 179–193 and resolves controlled numbering discrepancies.",
        "Run one past-project replay and one live-project soft launch; capture real hold durations, role conflicts and missing evidence before broad rollout.",
        "After policy sign-off, lock the final Gate authority, payment/release mapping, warranty triggers and enterprise ID/time-tracking integration.",
    ])
    doc.add_heading("Definition of done achieved in this release", level=2)
    done_rows = [
        ("Three layers connected", "Yes — 13 L1 nodes link to 14 L2 nodes; every L3 record has a valid L2 home/touchpoint."),
        ("All forms in L3", "Yes — 52/52 controlled records, including supplemental SMO."),
        ("Matrices in L3", "Yes — 57-row Approval Matrix and 18-row Responsibility Matrix."),
        ("Non-compliant client blocking", "Yes — authority, decision group, scale, fatal site, modular compatibility, route and engagement hard controls."),
        ("Form loophole blocking", "Yes — fields, applicability, N/A evidence, signature, approval, task and deletion protections."),
        ("UI/bug review", "Yes — responsive, accessible, clean-load and production-build checks passed."),
        ("No invented policy", "Yes — all explicit source TBDs remain visible for authorized decision."),
    ]
    add_table(doc, ["Acceptance criterion", "Result"], done_rows, widths=[2.1, 4.65], font_size=8.1, header_fill=GREEN)
    add_body(doc, "End of controlled report.", size=8, italic=True, color=MUTED)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    output = build_report()
    print(output)
